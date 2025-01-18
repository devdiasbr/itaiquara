from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import qrcode
from PIL import Image

# Gerar QR Code do PIX
qr = qrcode.QRCode(
    version=1,
    error_correction=qrcode.constants.ERROR_CORRECT_L,
    box_size=10,
    border=4,
)
qr.add_data('https://nubank.com.br/cobrar/1gw6o4/678a8abc-8861-4239-ae2e-f9026d010860')
qr.make(fit=True)
qr_img = qr.make_image(fill_color="black", back_color="white")
qr_img.save('pix_qr.png')

# Criar novo documento
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

# Título
title = doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS DE DESENVOLVIMENTO DE SOFTWARE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Partes
doc.add_paragraph('Por este instrumento particular, de um lado:')

doc.add_paragraph('CONTRATANTE: ITAIQUARA ALIMENTOS S.A. - EM RECUPERAÇÃO JUDICIAL, pessoa jurídica de direito privado, inscrita no CNPJ sob o nº 72.111.321/0001-74, com sede na Fazenda Itaiquara, Tapiratiba - SP, CEP 13.760-000, doravante denominada simplesmente CONTRATANTE;')

doc.add_paragraph('CONTRATADO: BRUNO HENRIQUE DE OLIVEIRA DIAS, pessoa física, desenvolvedor de software, inscrito no CPF sob nº 414.078.848-83, residente e domiciliado em [ENDEREÇO], doravante denominado simplesmente CONTRATADO.')

doc.add_paragraph('As partes acima identificadas têm, entre si, justo e acertado o presente Contrato de Prestação de Serviços de Desenvolvimento de Software, que se regerá pelas cláusulas seguintes e pelas condições descritas no presente.')

# Cláusulas
doc.add_heading('CLÁUSULA PRIMEIRA - DO OBJETO', 1)
doc.add_paragraph('O presente contrato tem como objeto a prestação de serviços de desenvolvimento de software pelo CONTRATADO, especificamente o desenvolvimento do sistema "KML ITAIQUARA 2.0", que consiste em uma aplicação para geração de arquivos KML com as seguintes funcionalidades principais:')
obj = doc.add_paragraph()
obj.add_run('\n• Interface gráfica para manipulação de dados')
obj.add_run('\n• Geração de arquivos KML')
obj.add_run('\n• Gerenciamento de dados em banco local')
obj.add_run('\n• Documentação completa do sistema')

doc.add_heading('CLÁUSULA SEGUNDA - DO PRAZO', 1)
doc.add_paragraph('O prazo para execução dos serviços será de 80 (oitenta) horas de desenvolvimento, com início após a aprovação e primeiro pagamento, conforme cronograma estabelecido entre as partes.')

doc.add_heading('CLÁUSULA TERCEIRA - DO VALOR E FORMA DE PAGAMENTO', 1)
doc.add_paragraph('O valor total dos serviços será de R$ 6.000,00 (seis mil reais), a serem pagos da seguinte forma:')
pag = doc.add_paragraph()
pag.add_run('\n• 40% (R$ 2.400,00) na aprovação e assinatura deste contrato')
pag.add_run('\n• 60% (R$ 3.600,00) na entrega final do projeto')

doc.add_paragraph('Parágrafo Único: Os pagamentos deverão ser realizados através de PIX para:')
dados_pix = doc.add_paragraph()
dados_pix.add_run('\nBanco: Nubank')
dados_pix.add_run('\nTitular: Bruno Henrique de Oliveira Dias')
dados_pix.add_run('\nCPF: 414.078.848-83')
dados_pix.add_run('\nConta: 888416687-8')
dados_pix.add_run('\nChaves PIX:')
dados_pix.add_run('\n  - CPF: 414.078.848-83')
dados_pix.add_run('\n  - Celular: (11) 95448-5244')

# Adicionar QR Code
doc.add_picture('pix_qr.png', width=Inches(2))

doc.add_heading('CLÁUSULA QUARTA - DAS OBRIGAÇÕES DO CONTRATADO', 1)
obrig = doc.add_paragraph()
obrig.add_run('O CONTRATADO se obriga a:')
obrig.add_run('\n\n• Desenvolver o software conforme especificações acordadas')
obrig.add_run('\n• Manter sigilo sobre todas as informações recebidas do CONTRATANTE')
obrig.add_run('\n• Entregar o código fonte completo e documentado')
obrig.add_run('\n• Fornecer suporte técnico durante 1 (um) mês após a entrega')
obrig.add_run('\n• Realizar o treinamento básico de uso do sistema')

doc.add_heading('CLÁUSULA QUINTA - DAS OBRIGAÇÕES DO CONTRATANTE', 1)
obrig_cont = doc.add_paragraph()
obrig_cont.add_run('O CONTRATANTE se obriga a:')
obrig_cont.add_run('\n\n• Fornecer todas as informações necessárias para o desenvolvimento')
obrig_cont.add_run('\n• Realizar os pagamentos nos prazos acordados')
obrig_cont.add_run('\n• Realizar a validação das entregas em tempo hábil')

doc.add_heading('CLÁUSULA SEXTA - DA PROPRIEDADE INTELECTUAL', 1)
doc.add_paragraph('Após a conclusão do serviço e quitação total dos valores, o CONTRATANTE terá direito de uso perpétuo do software desenvolvido, incluindo seu código fonte. O CONTRATADO manterá os direitos de propriedade intelectual sobre a metodologia e componentes genéricos desenvolvidos.')

doc.add_heading('CLÁUSULA SÉTIMA - DA MANUTENÇÃO', 1)
doc.add_paragraph('Após o período de 1 (um) mês de suporte incluso, caso seja do interesse do CONTRATANTE, poderá ser contratado o serviço de manutenção mensal no valor de R$ 700,00 (setecentos reais) que incluirá:')
manut = doc.add_paragraph()
manut.add_run('\n• Suporte em horário comercial (9h às 17h, segunda a sexta)')
manut.add_run('\n• Correções de bugs')
manut.add_run('\n• Atualizações de segurança')

doc.add_heading('CLÁUSULA OITAVA - DA RESCISÃO', 1)
doc.add_paragraph('Este contrato poderá ser rescindido por qualquer das partes, mediante comunicação expressa com antecedência mínima de 15 (quinze) dias, respeitando-se os serviços em andamento e os pagamentos proporcionais devidos.')

doc.add_heading('CLÁUSULA NONA - DO FORO', 1)
doc.add_paragraph('Fica eleito o foro da Comarca de [CIDADE] para dirimir quaisquer dúvidas ou controvérsias oriundas deste contrato, com renúncia expressa a qualquer outro, por mais privilegiado que seja.')

# Assinaturas
doc.add_paragraph('\n\nE por estarem assim justos e contratados, firmam o presente instrumento em duas vias de igual teor.')
doc.add_paragraph('\n\n[CIDADE], [DATA]')

doc.add_paragraph('\n\n_________________________________')
doc.add_paragraph('CONTRATANTE')
doc.add_paragraph('ITAIQUARA ALIMENTOS S.A. - EM RECUPERAÇÃO JUDICIAL')
doc.add_paragraph('CNPJ: 72.111.321/0001-74')

doc.add_paragraph('\n\n_________________________________')
doc.add_paragraph('CONTRATADO')
doc.add_paragraph('BRUNO HENRIQUE DE OLIVEIRA DIAS')
doc.add_paragraph('CPF: 414.078.848-83')

# Salvar o documento
doc.save('Contrato KML ITAIQUARA 2.0.docx')
